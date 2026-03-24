#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word (.docx) Report Generator for Regulatory Review
"""

from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─── Color palette ────────────────────────────────────────────────────────────
_C = {
    "header_bg":   RGBColor(0x1A, 0x3A, 0x5C),   # dark navy
    "header_fg":   RGBColor(0xFF, 0xFF, 0xFF),
    "section_bg":  RGBColor(0xE8, 0xF0, 0xFE),   # light blue
    "completed":   RGBColor(0x1E, 0x88, 0x55),
    "blocked":     RGBColor(0xC6, 0x28, 0x28),
    "in_progress": RGBColor(0xF5, 0x7F, 0x17),
    "pending":     RGBColor(0x75, 0x75, 0x75),
    "high_risk":   RGBColor(0xC6, 0x28, 0x28),
    "med_risk":    RGBColor(0xF5, 0x7F, 0x17),
    "low_risk":    RGBColor(0x1E, 0x88, 0x55),
}

STATUS_DISPLAY = {
    "completed":    "Completed",
    "in_progress":  "In Progress",
    "under_review": "Under Review",
    "blocked":      "Blocked",
    "pending":      "Pending",
}

RISK_DISPLAY = {
    "high":   "HIGH",
    "medium": "MEDIUM",
    "low":    "LOW",
}


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _rgb_to_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _add_heading(doc: "Document", text: str, level: int = 1):
    """Add styled heading."""
    para = doc.add_heading(text, level=level)
    para.runs[0].font.color.rgb = _C["header_bg"]
    return para


def generate_word_report(
    report: Dict,
    output_path: Optional[str] = None,
    project_path: Optional[Path] = None,
) -> Path:
    """Generate a professional .docx report from a review report dict."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Header banner ─────────────────────────────────────────────────────────
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = "Table Grid"
    left_cell  = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)

    _set_cell_bg(left_cell,  _rgb_to_hex(_C["header_bg"]))
    _set_cell_bg(right_cell, _rgb_to_hex(_C["header_bg"]))

    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = lp.add_run("REGULATORY REVIEW REPORT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = _C["header_fg"]

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = rp.add_run("UNIVERSAL INTEGRATED CORP.\nBD & Scientific Affairs")
    run2.font.size = Pt(9)
    run2.font.color.rgb = _C["header_fg"]

    doc.add_paragraph()

    # ── Project overview table ─────────────────────────────────────────────────
    _add_heading(doc, "Project Overview", level=1)

    ts = datetime.fromisoformat(report["review_date"]).strftime("%Y-%m-%d %H:%M")
    summary = report["summary"]
    rate_float = float(report["completion_rate"].replace("%", ""))

    overview_data = [
        ("Project",         report["project"].upper()),
        ("Document Type",   report.get("document_type", "—").replace("_", " ").title()),
        ("Review Date",     ts),
        ("Overall Status",  report["overall_status"].replace("_", " ").title()),
        ("Completion Rate", report["completion_rate"]),
        ("Completed Items", f"{summary['completed']} / {summary['total']}"),
        ("High-Risk Items", str(summary["high_risk_items"])),
    ]

    ov_table = doc.add_table(rows=len(overview_data), cols=2)
    ov_table.style = "Table Grid"
    ov_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (field, value) in enumerate(overview_data):
        label_cell = ov_table.cell(i, 0)
        value_cell = ov_table.cell(i, 1)

        _set_cell_bg(label_cell, _rgb_to_hex(_C["section_bg"]))

        lp = label_cell.paragraphs[0]
        lp.add_run(field).bold = True

        vp = value_cell.paragraphs[0]
        run = vp.add_run(value)

        # Color-code status and risk
        if field == "Overall Status":
            if "ready" in report["overall_status"]:
                run.font.color.rgb = _C["completed"]
            elif "attention" in report["overall_status"]:
                run.font.color.rgb = _C["blocked"]
            else:
                run.font.color.rgb = _C["in_progress"]
        elif field == "Completion Rate":
            run.bold = True
            if rate_float >= 70:
                run.font.color.rgb = _C["completed"]
            elif rate_float >= 40:
                run.font.color.rgb = _C["in_progress"]
            else:
                run.font.color.rgb = _C["blocked"]
        elif field == "High-Risk Items" and summary["high_risk_items"] > 0:
            run.font.color.rgb = _C["high_risk"]
            run.bold = True

    doc.add_paragraph()

    # ── Checklist table ────────────────────────────────────────────────────────
    _add_heading(doc, "Document Checklist", level=1)

    col_names = ["#", "Document Item", "Status", "Risk Level", "Notes"]
    col_widths = [Cm(1.0), Cm(6.5), Cm(3.0), Cm(2.5), Cm(5.5)]

    chk_table = doc.add_table(rows=1 + len(report["items"]), cols=5)
    chk_table.style = "Table Grid"

    # Header row
    for j, col_name in enumerate(col_names):
        cell = chk_table.cell(0, j)
        _set_cell_bg(cell, _rgb_to_hex(_C["header_bg"]))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col_name)
        run.bold = True
        run.font.color.rgb = _C["header_fg"]
        run.font.size = Pt(10)

    # Set column widths
    for j, width in enumerate(col_widths):
        for row in chk_table.rows:
            row.cells[j].width = width

    # Data rows
    for i, item in enumerate(report["items"], 1):
        row = chk_table.rows[i]

        # Alternating row shading
        if i % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, "F5F5F5")

        # #
        row.cells[0].paragraphs[0].add_run(str(i)).font.size = Pt(9)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Item name
        row.cells[1].paragraphs[0].add_run(item["item"]).font.size = Pt(9)

        # Status
        status_run = row.cells[2].paragraphs[0].add_run(
            STATUS_DISPLAY.get(item["status"], item["status"])
        )
        status_run.font.size = Pt(9)
        status_color = {
            "completed":    _C["completed"],
            "blocked":      _C["blocked"],
            "in_progress":  _C["in_progress"],
            "under_review": _C["in_progress"],
            "pending":      _C["pending"],
        }.get(item["status"])
        if status_color:
            status_run.font.color.rgb = status_color
        status_run.bold = item["status"] in ("completed", "blocked")

        # Risk
        risk_run = row.cells[3].paragraphs[0].add_run(
            RISK_DISPLAY.get(item["risk_level"], item["risk_level"])
        )
        risk_run.font.size = Pt(9)
        risk_color = {
            "high":   _C["high_risk"],
            "medium": _C["med_risk"],
            "low":    _C["low_risk"],
        }.get(item["risk_level"])
        if risk_color:
            risk_run.font.color.rgb = risk_color
        risk_run.bold = item["risk_level"] == "high"

        # Notes
        row.cells[4].paragraphs[0].add_run(item.get("notes") or "—").font.size = Pt(9)

    doc.add_paragraph()

    # ── Action items ──────────────────────────────────────────────────────────
    if report.get("action_items"):
        _add_heading(doc, "Action Items", level=1)

        action_table = doc.add_table(rows=1 + len(report["action_items"]), cols=3)
        action_table.style = "Table Grid"

        for j, hdr in enumerate(["Priority", "Document Item", "Required Action"]):
            cell = action_table.cell(0, j)
            _set_cell_bg(cell, _rgb_to_hex(_C["header_bg"]))
            p = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.bold = True
            run.font.color.rgb = _C["header_fg"]
            run.font.size = Pt(10)

        for i, action in enumerate(report["action_items"], 1):
            row = action_table.rows[i]
            if i % 2 == 0:
                for cell in row.cells:
                    _set_cell_bg(cell, "F5F5F5")

            priority_run = row.cells[0].paragraphs[0].add_run(action["priority"].upper())
            priority_run.bold = True
            priority_run.font.size = Pt(9)
            if action["priority"] == "high":
                priority_run.font.color.rgb = _C["high_risk"]
            else:
                priority_run.font.color.rgb = _C["med_risk"]

            row.cells[1].paragraphs[0].add_run(action["item"]).font.size = Pt(9)
            row.cells[2].paragraphs[0].add_run(action["action"]).font.size = Pt(9)

        doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Generated {ts}  |  UNIVERSAL INTEGRATED CORP.  |  BD & Scientific Affairs"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = _C["pending"]
    footer_run.italic = True

    # ── Save ──────────────────────────────────────────────────────────────────
    if output_path:
        dest = Path(output_path)
    elif project_path:
        dest = Path(project_path) / "review" / f"review-report-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
    else:
        dest = Path(f"review-report-{report['project']}-{datetime.now().strftime('%Y%m%d')}.docx")

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest
