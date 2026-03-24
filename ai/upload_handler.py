#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document upload and text extraction handler.
Supports PDF, Word (.docx), and Excel (.xlsx/.xls).
"""

from __future__ import annotations

import io
import math
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional


@dataclass
class ParsedDocument:
    filename: str
    file_type: str          # "pdf" | "docx" | "xlsx" | "unknown"
    raw_text: str
    chunks: List[str] = field(default_factory=list)
    page_count: int = 0
    char_count: int = 0
    error: Optional[str] = None


class DocumentParser:
    """Extract text from PDF, Word, and Excel files and split into chunks."""

    CHUNK_CHARS = 3000   # ~750 tokens at 4 chars/token
    CHUNK_OVERLAP = 300

    # ── Public API ─────────────────────────────────────────────────────────

    def parse(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """Detect file type from extension and extract text."""
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(file_bytes, filename)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file_bytes, filename)
        elif ext in (".xlsx", ".xls"):
            return self._parse_xlsx(file_bytes, filename)
        else:
            # Attempt to read as plain text (UTF-8 with fallback)
            try:
                text = file_bytes.decode("utf-8", errors="replace")
            except Exception as exc:
                return ParsedDocument(filename=filename, file_type="unknown",
                                      raw_text="", error=str(exc))
            doc = ParsedDocument(filename=filename, file_type="text",
                                 raw_text=text, char_count=len(text))
            doc.chunks = self.chunk_text(text)
            return doc

    # ── Parsers ────────────────────────────────────────────────────────────

    def _parse_pdf(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            import pypdf  # type: ignore
        except ImportError:
            try:
                import PyPDF2 as pypdf  # type: ignore  # legacy fallback
            except ImportError:
                return ParsedDocument(filename=filename, file_type="pdf",
                                      raw_text="",
                                      error="pypdf not installed. Run: pip install pypdf")

        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages: List[str] = []
            for page in reader.pages:
                text = page.extract_text() or ""
                pages.append(text)

            raw_text = "\n\n".join(pages)
            doc = ParsedDocument(
                filename=filename,
                file_type="pdf",
                raw_text=raw_text,
                page_count=len(pages),
                char_count=len(raw_text),
            )
            doc.chunks = self.chunk_text(raw_text)
            return doc
        except Exception as exc:
            return ParsedDocument(filename=filename, file_type="pdf",
                                  raw_text="", error=str(exc))

    def _parse_docx(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            from docx import Document  # python-docx
        except ImportError:
            return ParsedDocument(filename=filename, file_type="docx",
                                  raw_text="",
                                  error="python-docx not installed. Run: pip install python-docx")
        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract table content
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))

            raw_text = "\n\n".join(paragraphs)
            result = ParsedDocument(
                filename=filename,
                file_type="docx",
                raw_text=raw_text,
                page_count=0,
                char_count=len(raw_text),
            )
            result.chunks = self.chunk_text(raw_text)
            return result
        except Exception as exc:
            return ParsedDocument(filename=filename, file_type="docx",
                                  raw_text="", error=str(exc))

    def _parse_xlsx(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        try:
            import openpyxl  # type: ignore
        except ImportError:
            return ParsedDocument(filename=filename, file_type="xlsx",
                                  raw_text="",
                                  error="openpyxl not installed. Run: pip install openpyxl")
        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            sheets: List[str] = []
            for sheet in wb.worksheets:
                rows: List[str] = [f"=== Sheet: {sheet.title} ==="]
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        rows.append(" | ".join(cells))
                sheets.append("\n".join(rows))

            raw_text = "\n\n".join(sheets)
            result = ParsedDocument(
                filename=filename,
                file_type="xlsx",
                raw_text=raw_text,
                char_count=len(raw_text),
            )
            result.chunks = self.chunk_text(raw_text)
            return result
        except Exception as exc:
            return ParsedDocument(filename=filename, file_type="xlsx",
                                  raw_text="", error=str(exc))

    # ── Chunking ───────────────────────────────────────────────────────────

    def chunk_text(self, text: str,
                   chunk_size: int = CHUNK_CHARS,
                   overlap: int = CHUNK_OVERLAP) -> List[str]:
        """Split text into overlapping fixed-size character chunks."""
        if not text.strip():
            return []

        chunks: List[str] = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            # Try to break on sentence boundary
            if end < len(text):
                last_period = max(chunk.rfind("。"), chunk.rfind("."),
                                  chunk.rfind("\n"))
                if last_period > chunk_size // 2:
                    chunk = text[start: start + last_period + 1]
                    end = start + last_period + 1
            chunks.append(chunk.strip())
            start = end - overlap

        return [c for c in chunks if c]

    def estimate_tokens(self, text: str) -> int:
        """Rough token estimate: 1 token ≈ 4 English chars or 1.5 CJK chars."""
        cjk_chars = sum(1 for c in text if "\u4e00" <= c <= "\u9fff")
        other_chars = len(text) - cjk_chars
        return math.ceil(cjk_chars / 1.5 + other_chars / 4)
